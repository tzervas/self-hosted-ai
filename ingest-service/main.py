#!/usr/bin/env python3
"""
Document Ingest Automation Service

Automatically processes documents from watch directories:
- PDFs, Office docs, text files, images
- Generates embeddings for RAG
- Extracts metadata and creates summaries
- Integrates with Open WebUI knowledge base

Supports:
- Watch directories for automatic processing
- API for manual ingestion
- Batch processing
- OCR for images/scanned docs
- Multi-language support
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import os
import tempfile
import time
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional
from contextlib import asynccontextmanager

from fastapi import FastAPI, File, HTTPException, UploadFile, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import httpx
from watchdog.observers import Observer
from watchdog.observers.polling import PollingObserver
from watchdog.events import FileSystemEventHandler, FileCreatedEvent

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    MD = "markdown"
    HTML = "html"
    IMAGE = "image"
    CODE = "code"
    AUDIO = "audio"
    VIDEO = "video"
    UNKNOWN = "unknown"


class ProcessingStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


@dataclass
class ProcessedDocument:
    """Represents a processed document."""
    id: str
    filename: str
    filepath: str
    doc_type: DocumentType
    status: ProcessingStatus
    content_hash: str
    size_bytes: int
    created_at: datetime
    processed_at: Optional[datetime] = None
    chunks: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error_message: Optional[str] = None


class IngestRequest(BaseModel):
    filepath: str = Field(..., description="Path to the document to ingest")
    collection: str = Field(default="default", description="Target collection")
    tags: List[str] = Field(default_factory=list, description="Tags to apply")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Additional metadata")


class IngestResponse(BaseModel):
    success: bool
    document_id: str
    message: str
    chunks_created: int = 0


class IngestConfig(BaseModel):
    watch_directories: List[str] = Field(default_factory=list)
    supported_extensions: List[str] = Field(
        default=[".pdf", ".docx", ".doc", ".txt", ".md", ".html", ".py", ".js", ".ts", ".json", ".yaml", ".yml"]
    )
    chunk_size: int = Field(default=1000, description="Characters per chunk")
    chunk_overlap: int = Field(default=200, description="Overlap between chunks")
    embedding_model: str = Field(default="nomic-embed-text")
    ollama_url: str = Field(default="http://ollama-cpu:11434")
    openwebui_url: str = Field(default="http://open-webui:8080")
    max_file_size_mb: int = Field(default=100)


class DocumentProcessor:
    """Processes various document types."""
    
    EXTENSION_TO_TYPE = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOC,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MD,
        ".html": DocumentType.HTML,
        ".htm": DocumentType.HTML,
        ".py": DocumentType.CODE,
        ".js": DocumentType.CODE,
        ".ts": DocumentType.CODE,
        ".java": DocumentType.CODE,
        ".cpp": DocumentType.CODE,
        ".c": DocumentType.CODE,
        ".go": DocumentType.CODE,
        ".rs": DocumentType.CODE,
        ".json": DocumentType.CODE,
        ".yaml": DocumentType.CODE,
        ".yml": DocumentType.CODE,
        ".png": DocumentType.IMAGE,
        ".jpg": DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".gif": DocumentType.IMAGE,
        ".webp": DocumentType.IMAGE,
        ".mp3": DocumentType.AUDIO,
        ".wav": DocumentType.AUDIO,
        ".m4a": DocumentType.AUDIO,
        ".mp4": DocumentType.VIDEO,
        ".webm": DocumentType.VIDEO,
        ".mkv": DocumentType.VIDEO,
    }
    
    def __init__(self, config: IngestConfig):
        self.config = config
        self.documents: Dict[str, ProcessedDocument] = {}
        
    def detect_type(self, filepath: str) -> DocumentType:
        """Detect document type from extension."""
        ext = Path(filepath).suffix.lower()
        return self.EXTENSION_TO_TYPE.get(ext, DocumentType.UNKNOWN)
    
    def compute_hash(self, filepath: str) -> str:
        """Compute content hash for deduplication."""
        hasher = hashlib.sha256()
        with open(filepath, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()[:16]
    
    async def extract_text(self, filepath: str, doc_type: DocumentType) -> str:
        """Extract text content from document."""
        try:
            if doc_type == DocumentType.TXT or doc_type == DocumentType.MD:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
            elif doc_type == DocumentType.CODE:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
            elif doc_type == DocumentType.PDF:
                return await self._extract_pdf(filepath)
                
            elif doc_type in (DocumentType.DOCX, DocumentType.DOC):
                return await self._extract_docx(filepath)
                
            elif doc_type == DocumentType.HTML:
                return await self._extract_html(filepath)
                
            elif doc_type == DocumentType.IMAGE:
                return await self._extract_image_ocr(filepath)
                
            elif doc_type == DocumentType.AUDIO:
                return await self._transcribe_audio(filepath)
                
            else:
                # Try to read as text
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {e}")
            return ""
    
    async def _extract_pdf(self, filepath: str) -> str:
        """Extract text from PDF using pypdf or pdfplumber."""
        try:
            # Try pypdf first
            import pypdf
            reader = pypdf.PdfReader(filepath)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            try:
                # Fallback to pdfplumber
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    return "\n".join(page.extract_text() or "" for page in pdf.pages)
            except ImportError:
                logger.warning("No PDF library available (install pypdf or pdfplumber)")
                return ""
    
    async def _extract_docx(self, filepath: str) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed")
            return ""
    
    async def _extract_html(self, filepath: str) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text(separator='\n', strip=True)
        except ImportError:
            # Fallback: simple regex
            import re
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
            text = re.sub(r'<[^>]+>', '', text)
            return text
    
    async def _extract_image_ocr(self, filepath: str) -> str:
        """Extract text from images using OCR."""
        try:
            # Use ollama vision model for OCR
            async with httpx.AsyncClient() as client:
                import base64
                with open(filepath, 'rb') as f:
                    image_data = base64.b64encode(f.read()).decode()
                
                response = await client.post(
                    f"{self.config.ollama_url}/api/generate",
                    json={
                        "model": "llava",
                        "prompt": "Extract all text from this image. Return only the extracted text, nothing else.",
                        "images": [image_data],
                        "stream": False
                    },
                    timeout=60.0
                )
                
                if response.status_code == 200:
                    return response.json().get("response", "")
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
        return ""
    
    async def _transcribe_audio(self, filepath: str) -> str:
        """Transcribe audio using Whisper."""
        try:
            # Check for local whisper service first
            async with httpx.AsyncClient() as client:
                with open(filepath, 'rb') as f:
                    response = await client.post(
                        "http://whisper:9000/transcribe",
                        files={"file": f},
                        timeout=300.0
                    )
                if response.status_code == 200:
                    return response.json().get("text", "")
        except Exception as e:
            logger.warning(f"Audio transcription failed: {e}")
        return ""
    
    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks."""
        chunks = []
        
        if not text:
            return chunks
        
        chunk_size = self.config.chunk_size
        overlap = self.config.chunk_overlap
        
        # Split by paragraphs first for better context
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append({
                        "text": current_chunk.strip(),
                        "index": len(chunks),
                        "char_count": len(current_chunk)
                    })
                
                # Start new chunk with overlap
                if len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + para + "\n\n"
                else:
                    current_chunk = para + "\n\n"
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                "text": current_chunk.strip(),
                "index": len(chunks),
                "char_count": len(current_chunk)
            })
        
        return chunks
    
    async def generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Generate embeddings for chunks using Ollama."""
        try:
            async with httpx.AsyncClient() as client:
                for chunk in chunks:
                    response = await client.post(
                        f"{self.config.ollama_url}/api/embeddings",
                        json={
                            "model": self.config.embedding_model,
                            "prompt": chunk["text"]
                        },
                        timeout=30.0
                    )
                    
                    if response.status_code == 200:
                        chunk["embedding"] = response.json().get("embedding")
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
        
        return chunks
    
    async def process_document(
        self,
        filepath: str,
        collection: str = "default",
        tags: List[str] = None,
        metadata: Dict[str, Any] = None
    ) -> ProcessedDocument:
        """Process a single document."""
        tags = tags or []
        metadata = metadata or {}
        
        # Create document record
        doc_id = f"doc_{int(time.time())}_{hashlib.md5(filepath.encode()).hexdigest()[:8]}"
        doc_type = self.detect_type(filepath)
        content_hash = self.compute_hash(filepath)
        
        doc = ProcessedDocument(
            id=doc_id,
            filename=os.path.basename(filepath),
            filepath=filepath,
            doc_type=doc_type,
            status=ProcessingStatus.PROCESSING,
            content_hash=content_hash,
            size_bytes=os.path.getsize(filepath),
            created_at=datetime.now(),
            metadata={
                "collection": collection,
                "tags": tags,
                **metadata
            }
        )
        
        self.documents[doc_id] = doc
        
        try:
            # Extract text
            text = await self.extract_text(filepath, doc_type)
            
            if not text:
                doc.status = ProcessingStatus.FAILED
                doc.error_message = "No text content extracted"
                return doc
            
            # Chunk text
            chunks = self.chunk_text(text)
            
            # Generate embeddings
            chunks = await self.generate_embeddings(chunks)
            
            doc.chunks = chunks
            doc.status = ProcessingStatus.COMPLETED
            doc.processed_at = datetime.now()
            
            logger.info(f"Processed {filepath}: {len(chunks)} chunks")
            
        except Exception as e:
            doc.status = ProcessingStatus.FAILED
            doc.error_message = str(e)
            logger.error(f"Failed to process {filepath}: {e}")
        
        return doc


class DirectoryWatcher(FileSystemEventHandler):
    """Watch directories for new files."""
    
    def __init__(self, processor: DocumentProcessor, supported_extensions: List[str]):
        self.processor = processor
        self.supported_extensions = supported_extensions
        self.processing_queue = asyncio.Queue()
    
    def on_created(self, event):
        if isinstance(event, FileCreatedEvent):
            ext = Path(event.src_path).suffix.lower()
            if ext in self.supported_extensions:
                logger.info(f"New file detected: {event.src_path}")
                asyncio.create_task(self._queue_file(event.src_path))
    
    async def _queue_file(self, filepath: str):
        await self.processing_queue.put(filepath)
    
    async def process_queue(self):
        """Process files from queue."""
        while True:
            try:
                filepath = await asyncio.wait_for(
                    self.processing_queue.get(),
                    timeout=1.0
                )
                
                # Wait a moment for file to finish writing
                await asyncio.sleep(2)
                
                if os.path.exists(filepath):
                    await self.processor.process_document(filepath)
                    
            except asyncio.TimeoutError:
                continue
            except Exception as e:
                logger.error(f"Queue processing error: {e}")


# Global state
config = IngestConfig(
    watch_directories=[
        os.environ.get("WATCH_DIR", "/data/ingest"),
        os.environ.get("WATCH_DIR_2", "/data/documents")
    ],
    ollama_url=os.environ.get("OLLAMA_URL", "http://ollama-cpu:11434"),
    openwebui_url=os.environ.get("OPENWEBUI_URL", "http://open-webui:8080"),
)

processor = DocumentProcessor(config)
watcher = DirectoryWatcher(processor, config.supported_extensions)

# Use PollingObserver as fallback when inotify limits are exceeded
# Set USE_POLLING=true env var or it will auto-fallback on inotify errors
USE_POLLING = os.environ.get("USE_POLLING", "false").lower() == "true"

def create_observer():
    """Create file system observer with automatic fallback to polling."""
    if USE_POLLING:
        logger.info("Using PollingObserver (explicit configuration)")
        return PollingObserver(timeout=5)
    try:
        obs = Observer()
        # Test if inotify works by checking watch limit
        import subprocess
        result = subprocess.run(
            ["cat", "/proc/sys/fs/inotify/max_user_watches"],
            capture_output=True, text=True
        )
        max_watches = int(result.stdout.strip()) if result.returncode == 0 else 0
        if max_watches < 8192:
            logger.warning(f"Low inotify limit ({max_watches}), using PollingObserver")
            return PollingObserver(timeout=5)
        return obs
    except Exception as e:
        logger.warning(f"inotify check failed ({e}), using PollingObserver")
        return PollingObserver(timeout=5)

observer = create_observer()


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan management."""
    # Start directory watchers
    for watch_dir in config.watch_directories:
        if os.path.exists(watch_dir):
            observer.schedule(watcher, watch_dir, recursive=True)
            logger.info(f"Watching directory: {watch_dir}")
    
    observer.start()
    
    # Start queue processor
    queue_task = asyncio.create_task(watcher.process_queue())
    
    yield
    
    # Cleanup
    queue_task.cancel()
    observer.stop()
    observer.join()


app = FastAPI(
    title="Document Ingest Service",
    description="Automated document processing and RAG integration",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health():
    """Health check."""
    return {
        "status": "healthy",
        "watched_directories": config.watch_directories,
        "documents_processed": len(processor.documents)
    }


@app.get("/status")
async def status():
    """Get service status."""
    return {
        "config": config.dict(),
        "documents": {
            "total": len(processor.documents),
            "completed": sum(1 for d in processor.documents.values() if d.status == ProcessingStatus.COMPLETED),
            "failed": sum(1 for d in processor.documents.values() if d.status == ProcessingStatus.FAILED),
            "processing": sum(1 for d in processor.documents.values() if d.status == ProcessingStatus.PROCESSING),
        }
    }


@app.post("/ingest", response_model=IngestResponse)
async def ingest_document(request: IngestRequest, background_tasks: BackgroundTasks):
    """Ingest a document."""
    if not os.path.exists(request.filepath):
        raise HTTPException(status_code=404, detail=f"File not found: {request.filepath}")
    
    doc = await processor.process_document(
        request.filepath,
        collection=request.collection,
        tags=request.tags,
        metadata=request.metadata
    )
    
    return IngestResponse(
        success=doc.status == ProcessingStatus.COMPLETED,
        document_id=doc.id,
        message=doc.error_message or "Document processed successfully",
        chunks_created=len(doc.chunks)
    )


@app.post("/ingest/upload")
async def upload_and_ingest(
    file: UploadFile = File(...),
    collection: str = "default"
):
    """Upload and ingest a file."""
    # Save to temp file
    suffix = Path(file.filename).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
    
    try:
        doc = await processor.process_document(tmp_path, collection=collection)
        return IngestResponse(
            success=doc.status == ProcessingStatus.COMPLETED,
            document_id=doc.id,
            message=doc.error_message or "Document processed successfully",
            chunks_created=len(doc.chunks)
        )
    finally:
        os.unlink(tmp_path)


@app.get("/documents")
async def list_documents(limit: int = 100, offset: int = 0):
    """List processed documents."""
    docs = list(processor.documents.values())[offset:offset + limit]
    return {
        "total": len(processor.documents),
        "documents": [
            {
                "id": d.id,
                "filename": d.filename,
                "type": d.doc_type.value,
                "status": d.status.value,
                "chunks": len(d.chunks),
                "created_at": d.created_at.isoformat(),
                "processed_at": d.processed_at.isoformat() if d.processed_at else None
            }
            for d in docs
        ]
    }


@app.get("/documents/{doc_id}")
async def get_document(doc_id: str):
    """Get document details."""
    if doc_id not in processor.documents:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc = processor.documents[doc_id]
    return {
        "id": doc.id,
        "filename": doc.filename,
        "filepath": doc.filepath,
        "type": doc.doc_type.value,
        "status": doc.status.value,
        "content_hash": doc.content_hash,
        "size_bytes": doc.size_bytes,
        "chunks": len(doc.chunks),
        "metadata": doc.metadata,
        "created_at": doc.created_at.isoformat(),
        "processed_at": doc.processed_at.isoformat() if doc.processed_at else None,
        "error": doc.error_message
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8200)
